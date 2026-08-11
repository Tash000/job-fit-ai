"""
Cover-letter exporters.

TXT    — plain text
LaTeX  — a professional LaTeX letter (the formatting source of truth)
PDF    — LaTeX compiled with a TeX engine (pdflatex/xelatex via MiKTeX/TeX Live)
DOCX   — converted from the SAME LaTeX with pandoc, falling back to python-docx

Tool discovery order: environment variable override → PATH → known install dirs.
Set PANDOC_PATH / PDFLATEX_PATH (or XELATEX_PATH) explicitly if they are not on PATH.
"""

import io
import os
import re
import shutil
import subprocess
import tempfile
from typing import Dict, Optional

try:
    import docx as _docx
    from docx.shared import Inches as _Inches, Pt as _Pt
except ImportError:  # pragma: no cover
    _docx = None

# ── Tool discovery ─────────────────────────────────────────────────────────────

# Known install locations on this machine (MiKTeX + Pandoc are portable installs).
_KNOWN_TOOL_DIRS = [
    r"D:\Downloaded Program Files\Pandoc",
    r"D:\Downloaded Program Files\MiKTeX\miktex\bin\x64",
    r"D:\Downloaded\miktex\bin\x64",
]

_ENGINE_CANDIDATES = ["pdflatex", "xelatex", "lualatex"]


def _find_tool(name: str, env_var: str = "") -> Optional[str]:
    """Locate an executable: env override → PATH → known install dirs."""
    if env_var:
        candidate = os.environ.get(env_var)
        if candidate and os.path.isfile(candidate):
            return candidate
    found = shutil.which(name)
    if found:
        return found
    exe = name + (".exe" if os.name == "nt" else "")
    for directory in _KNOWN_TOOL_DIRS:
        candidate = os.path.join(directory, exe)
        if os.path.isfile(candidate):
            return candidate
    return None


def find_pdf_engine() -> Optional[str]:
    for name in _ENGINE_CANDIDATES:
        path = _find_tool(name, f"{name.upper()}_PATH")
        if path:
            return path
    return None


def find_pandoc() -> Optional[str]:
    return _find_tool("pandoc", "PANDOC_PATH")


# ── LaTeX escaping ─────────────────────────────────────────────────────────────

_LATEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
    '"': "''",  # double quotes → typographic pair
}


def _latex_escape(text: str) -> str:
    out = text
    for char, esc in _LATEX_ESCAPES.items():
        out = out.replace(char, esc)
    return out


def _split_letter(text: str) -> Dict[str, str]:
    """Split a generated letter into salutation / body / closing lines."""
    paragraphs = [p.strip() for p in (text or "").split("\n\n") if p.strip()]
    salutation = ""
    closing = "Sincerely,"

    if paragraphs and re.match(r"^(Dear|Sehr geehrte|Hallo)", paragraphs[0]):
        salutation = paragraphs.pop(0)

    if paragraphs:
        last = paragraphs[-1]
        if re.match(r"^(Sincerely|Kind regards|Best regards|With regards|Mit freundlichen Gr)", last, re.IGNORECASE):
            closing = last
            paragraphs.pop()

    return {"salutation": salutation, "body": "\n\n".join(paragraphs), "closing": closing}


# ── Exporters ──────────────────────────────────────────────────────────────────


def export_txt(cover_letter_text: str) -> bytes:
    """Export cover letter as plain text bytes."""
    return (cover_letter_text or "").encode("utf-8")


def export_latex(
    cover_letter_text: str,
    candidate_name: str = "Candidate Name",
    candidate_email: str = "email@example.com",
    candidate_phone: str = "",
    candidate_address: str = "",
    candidate_links: str = "",
    company: str = "",
    position: str = "",
    subject: str = "",
) -> bytes:
    """Export the cover letter as a professional LaTeX letter (DIN 5008-style)."""
    parts = _split_letter(cover_letter_text)
    salutation = parts["salutation"] or "Dear Hiring Team,"
    body_paras = [p for p in parts["body"].split("\n\n") if p.strip()]
    closing = parts["closing"]

    esc = _latex_escape
    contact_lines = [f"\\textbf{{{esc(candidate_name)}}}"]
    if candidate_address:
        contact_lines.append(esc(candidate_address))
    if candidate_email:
        contact_lines.append(f"Email: \\href{{mailto:{candidate_email}}}{{{esc(candidate_email)}}}")
    if candidate_phone:
        contact_lines.append(f"Phone: {esc(candidate_phone)}")
    if candidate_links:
        contact_lines.append(esc(candidate_links))

    recipient = esc(company or "Hiring Team")
    subject_line = esc(subject or (f"Application for {position or 'the position'} — {company or 'your team'}"))
    body_tex = "\n\n".join(esc(p) for p in body_paras)
    address_block = " \\\\".join(contact_lines)  # LaTeX line break between address lines

    # Raw template — LaTeX commands are literal; letter content injected via
    # __PLACEHOLDERS__ that never occur in letter text.
    latex_template = r"""\documentclass[11pt,a4paper]{letter}
\usepackage[T1]{fontenc}
\usepackage[utf8]{inputenc}
\usepackage[left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm]{geometry}
\usepackage{parskip}
\usepackage[hidelinks]{hyperref}
\pagestyle{empty}

\date{\today}

\address{
    __ADDRESS__
}

\signature{__NAME__}
\name{__NAME__}
\begin{document}

\begin{letter}{__RECIPIENT__}

\textbf{__SUBJECT__}\par\vspace{0.6em}
\opening{__SALUTATION__}
__BODY__

\closing{__CLOSING__}

\encl{CV / Resume \\ Transcripts \\ Certificates \\ Portfolio links}

\end{letter}
\end{document}
"""

    latex = (
        latex_template
        .replace("__ADDRESS__", address_block)
        .replace("__NAME__", esc(candidate_name))
        .replace("__RECIPIENT__", recipient)
        .replace("__SUBJECT__", subject_line)
        .replace("__SALUTATION__", esc(salutation))
        .replace("__BODY__", body_tex)
        .replace("__CLOSING__", esc(closing))
    )
    return latex.encode("utf-8")


def _compile_pdf(latex_bytes: bytes) -> Optional[bytes]:
    """Compile LaTeX to PDF bytes. Returns None when no engine is available."""
    engine = find_pdf_engine()
    if not engine:
        return None
    with tempfile.TemporaryDirectory(prefix="vitralume_cl_") as tmp:
        tex_path = os.path.join(tmp, "cover_letter.tex")
        with open(tex_path, "wb") as fh:
            fh.write(latex_bytes)
        try:
            proc = subprocess.run(
                [engine, "-interaction=nonstopmode", "-halt-on-error", "cover_letter.tex"],
                cwd=tmp,
                capture_output=True,
                timeout=120,
            )
            if proc.returncode != 0:
                return None
            pdf_path = os.path.join(tmp, "cover_letter.pdf")
            if not os.path.isfile(pdf_path):
                return None
            with open(pdf_path, "rb") as fh:
                return fh.read()
        except (subprocess.SubprocessError, OSError):
            return None


def export_pdf(
    cover_letter_text: str,
    candidate_name: str = "Candidate Name",
    candidate_email: str = "email@example.com",
    candidate_phone: str = "",
    candidate_address: str = "",
    candidate_links: str = "",
    company: str = "",
    position: str = "",
    subject: str = "",
) -> bytes:
    """Compile the LaTeX letter into a PDF. Raises RuntimeError if no TeX engine is available."""
    latex_bytes = export_latex(
        cover_letter_text, candidate_name, candidate_email, candidate_phone,
        candidate_address, candidate_links, company, position, subject,
    )
    pdf = _compile_pdf(latex_bytes)
    if pdf is None:
        raise RuntimeError(
            "PDF export needs a LaTeX engine (pdflatex/xelatex). "
            "Install MiKTeX or TeX Live, or set the PDFLATEX_PATH environment variable."
        )
    return pdf


def export_docx(
    cover_letter_text: str,
    candidate_name: str = "Candidate Name",
    candidate_email: str = "email@example.com",
    candidate_phone: str = "",
    candidate_address: str = "",
    candidate_links: str = "",
    company: str = "",
    position: str = "",
    subject: str = "",
) -> bytes:
    """
    Export a Word document. Prefers converting the SAME LaTeX source with pandoc
    (formatting stays identical to the PDF); falls back to python-docx, then plain text.
    """
    latex_bytes = export_latex(
        cover_letter_text, candidate_name, candidate_email, candidate_phone,
        candidate_address, candidate_links, company, position, subject,
    )

    pandoc = find_pandoc()
    if pandoc:
        try:
            with tempfile.TemporaryDirectory(prefix="vitralume_cl_") as tmp:
                tex_path = os.path.join(tmp, "cover_letter.tex")
                out_path = os.path.join(tmp, "cover_letter.docx")
                with open(tex_path, "wb") as fh:
                    fh.write(latex_bytes)
                proc = subprocess.run(
                    [pandoc, tex_path, "-o", out_path],
                    capture_output=True,
                    timeout=120,
                )
                if proc.returncode == 0 and os.path.isfile(out_path):
                    with open(out_path, "rb") as fh:
                        return fh.read()
        except (subprocess.SubprocessError, OSError):
            pass  # fall through to python-docx

    if _docx is not None:
        return _docx_fallback(cover_letter_text, candidate_name, candidate_email, candidate_phone)

    raise RuntimeError(
        "DOCX export needs pandoc or the python-docx package. "
        "Install pandoc (PANDOC_PATH env var) or run: pip install python-docx"
    )


def _docx_fallback(
    cover_letter_text: str,
    candidate_name: str,
    candidate_email: str,
    candidate_phone: str,
) -> bytes:
    """python-docx based DOCX generation (no pandoc required)."""
    doc = _docx.Document()

    for section in doc.sections:
        section.top_margin = _Inches(1.0)
        section.bottom_margin = _Inches(1.0)
        section.left_margin = _Inches(1.0)
        section.right_margin = _Inches(1.0)

    p_header = doc.add_paragraph()
    p_header.paragraph_format.space_after = _Pt(24)
    run_name = p_header.add_run(f"{candidate_name}\n")
    run_name.bold = True
    run_name.font.size = _Pt(14)
    run_name.font.name = "Arial"

    contact = " | ".join(filter(None, [f"Email: {candidate_email}", f"Phone: {candidate_phone}"]))
    if contact:
        run_contact = p_header.add_run(contact)
        run_contact.font.size = _Pt(10)
        run_contact.font.name = "Arial"
        run_contact.font.italic = True

    for p_text in (cover_letter_text or "").split("\n\n"):
        if not p_text.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = _Pt(12)
        run = p.add_run(p_text.strip())
        run.font.name = "Arial"
        run.font.size = _Pt(11)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
